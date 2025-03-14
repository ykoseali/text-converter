#!/usr/bin/env python3
"""
All-in-One Document Converter with Dependency Fixing

This script converts various document formats to text files and automatically
handles dependency installation when needed.

Features:
- Automatic dependency installation
- Fallback to direct extraction methods when needed
- Support for PDF, Word, Excel, CSV, XML and other formats
- Works even when dependencies are missing

Made by Yigit KOSEALI
"""

import os
import argparse
from pathlib import Path
import sys
import time
import re
import csv
import io
import subprocess
import tempfile
import platform
import shutil
from contextlib import contextmanager


def print_section(title):
    """Print a section title"""
    print("\n" + "=" * 60)
    print(f" {title}")
    print("=" * 60)


def install_dependency(package, message=None):
    """Install a Python package if not already installed"""
    if message:
        print(message)
    
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        print(f"Successfully installed {package}")
        return True
    except Exception as e:
        print(f"Failed to install {package}: {str(e)}")
        return False


def check_and_install_dependencies(required_formats):
    """Check and install necessary dependencies based on required formats"""
    print_section("Checking dependencies")
    
    installed_packages = []
    
    # Check for pandas (needed for Excel and CSV)
    if 'excel' in required_formats or 'csv' in required_formats:
        try:
            import pandas
            print("✓ pandas is already installed")
        except ImportError:
            if install_dependency("pandas", "Installing pandas for Excel/CSV support..."):
                installed_packages.append("pandas")
                # Also install Excel support
                install_dependency("openpyxl", "Installing openpyxl for modern Excel files...")
                install_dependency("xlrd", "Installing xlrd for old Excel files...")
                installed_packages.extend(["openpyxl", "xlrd"])
    
    # Check for docx support
    if 'docx' in required_formats:
        try:
            import docx
            print("✓ python-docx is already installed")
        except ImportError:
            if install_dependency("python-docx", "Installing python-docx for DOCX support..."):
                installed_packages.append("python-docx")
    
    # Check for PDF support
    if 'pdf' in required_formats:
        try:
            import PyPDF2
            print("✓ PyPDF2 is already installed")
        except ImportError:
            if install_dependency("PyPDF2", "Installing PyPDF2 for PDF support..."):
                installed_packages.append("PyPDF2")
    
    # Try to install unstructured with basic dependencies
    try:
        import unstructured
        print(f"✓ unstructured is already installed (version {unstructured.__version__})")
    except (ImportError, AttributeError):
        print("Installing unstructured for advanced document processing...")
        if install_dependency("unstructured"):
            installed_packages.append("unstructured")
    
    # Install NLTK if needed for unstructured
    try:
        import nltk
        
        # Check if punkt is already downloaded
        try:
            nltk.data.find('tokenizers/punkt')
            print("✓ NLTK punkt is already installed")
        except LookupError:
            print("Downloading NLTK punkt...")
            nltk.download('punkt')
        
        # Check if averaged_perceptron_tagger is already downloaded
        try:
            nltk.data.find('taggers/averaged_perceptron_tagger')
            print("✓ NLTK tagger is already installed")
        except LookupError:
            print("Downloading NLTK averaged_perceptron_tagger...")
            nltk.download('averaged_perceptron_tagger')
    
    except ImportError:
        if install_dependency("nltk", "Installing NLTK for text processing..."):
            installed_packages.append("nltk")
            # Download needed NLTK data
            try:
                import nltk
                print("Downloading NLTK data...")
                nltk.download('punkt')
                nltk.download('averaged_perceptron_tagger')
            except Exception as e:
                print(f"Failed to download NLTK data: {str(e)}")
    
    # Check system dependencies
    if platform.system() == 'Darwin':  # macOS
        check_macos_dependencies(required_formats)
    
    # If new packages were installed, recommend restarting
    if installed_packages:
        print("\nNew packages were installed. It's recommended to restart the script")
        print("for changes to take effect. Would you like to restart now?")
        response = input("Restart? (y/n): ")
        if response.lower() in ['y', 'yes']:
            print("\nRestarting script...\n")
            os.execv(sys.executable, [sys.executable] + sys.argv)
    
    return True


def check_macos_dependencies(required_formats):
    """Check and report on macOS system dependencies"""
    print("\nChecking macOS system dependencies:")
    
    # Check if homebrew is installed
    if not shutil.which('brew'):
        print("- Homebrew is not installed. Many dependencies require it.")
        print("  Install with: /bin/bash -c \"$(curl -fsSL https://raw.githubusercontent.com/Homebrew/install/HEAD/install.sh)\"")
        return
    
    # Check needed tools
    dependencies = []
    
    if 'pdf' in required_formats:
        dependencies.append(('poppler', 'brew install poppler', 'for better PDF processing'))
    
    dependencies.extend([
        ('tesseract', 'brew install tesseract', 'for OCR/image processing'),
        ('textutil', '', 'for DOC files (built into macOS)'), 
        ('libreoffice', 'brew install libreoffice', 'for old Office formats')
    ])
    
    for dep, install_cmd, purpose in dependencies:
        if shutil.which(dep):
            print(f"✓ {dep} is installed {purpose}")
        else:
            if install_cmd:
                print(f"✗ {dep} is not installed {purpose}")
                print(f"  Install with: {install_cmd}")
            else:
                print(f"✗ {dep} is not installed {purpose}")


def extract_text_from_doc(file_path):
    """Extract text from a DOC file using multiple fallback methods"""
    
    # Method 1: Try using textutil on macOS (most reliable for macOS)
    if platform.system() == 'Darwin':  # macOS
        try:
            temp_txt = tempfile.NamedTemporaryFile(delete=False, suffix='.txt')
            temp_txt.close()
            
            result = subprocess.run([
                'textutil', '-convert', 'txt', '-output', 
                temp_txt.name, str(file_path)
            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            if result.returncode == 0:
                with open(temp_txt.name, 'r', errors='ignore') as f:
                    text = f.read()
                os.unlink(temp_txt.name)
                return text
        except Exception as e:
            print(f"  macOS textutil failed: {str(e)}")
    
    # Method 2: Try using textract if available
    try:
        import textract
        text = textract.process(file_path).decode('utf-8')
        return text
    except ImportError:
        print("  textract not installed, trying alternative methods...")
    except Exception as e:
        print(f"  textract extraction failed: {str(e)}")
    
    # Method 3: Try using antiword if available
    try:
        result = subprocess.run(['antiword', str(file_path)], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode == 0:
            return result.stdout.decode('utf-8')
    except FileNotFoundError:
        print("  antiword not found, trying next method...")
    except Exception as e:
        print(f"  antiword extraction failed: {str(e)}")
    
    # Method 4: Try using catdoc if available
    try:
        result = subprocess.run(['catdoc', str(file_path)], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode == 0:
            return result.stdout.decode('utf-8')
    except FileNotFoundError:
        print("  catdoc not found, trying next method...")
    except Exception as e:
        print(f"  catdoc extraction failed: {str(e)}")
    
    # Method 5: Try using libreoffice to convert to txt
    try:
        temp_dir = tempfile.mkdtemp()
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'txt:Text', 
            '--outdir', temp_dir, str(file_path)
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        if result.returncode == 0:
            txt_filename = os.path.basename(str(file_path)).replace('.doc', '.txt')
            txt_path = os.path.join(temp_dir, txt_filename)
            with open(txt_path, 'r', errors='ignore') as f:
                text = f.read()
            return text
    except FileNotFoundError:
        print("  libreoffice not found...")
    except Exception as e:
        print(f"  libreoffice conversion failed: {str(e)}")
    
    # Last resort: try to read the file directly
    try:
        with open(file_path, 'rb') as f:
            data = f.read()
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = data.decode(encoding, errors='ignore')
                    # Look for readable text content
                    if re.search(r'[A-Za-z]{5,}', text):
                        return text
                except UnicodeDecodeError:
                    continue
    except Exception:
        pass
    
    # If we get here, all methods failed
    raise Exception("Failed to extract text from DOC file. Try installing textract, antiword, catdoc, or LibreOffice.")


def extract_text_from_docx(file_path):
    """Extract text from a DOCX file directly"""
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except ImportError:
        # Try to install python-docx
        if install_dependency("python-docx", "Installing python-docx for DOCX support..."):
            # Try again after installation
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        else:
            raise Exception("Failed to install python-docx. Please install it manually: pip install python-docx")
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file directly"""
    try:
        import PyPDF2
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                extracted_text = page.extract_text() or ""
                text.append(extracted_text)
        return '\n\n'.join(text)
    except ImportError:
        # Try to install PyPDF2
        if install_dependency("PyPDF2", "Installing PyPDF2 for PDF support..."):
            # Try again after installation
            import PyPDF2
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    extracted_text = page.extract_text() or ""
                    text.append(extracted_text)
            return '\n\n'.join(text)
        else:
            # Try pdftotext if PyPDF2 installation fails
            try:
                with tempfile.NamedTemporaryFile(suffix='.txt') as temp_txt:
                    subprocess.run(
                        ['pdftotext', str(file_path), temp_txt.name],
                        check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
                    )
                    with open(temp_txt.name, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read()
            except Exception as pdf_e:
                raise Exception(f"Failed to extract text from PDF: {str(pdf_e)}")
    except Exception as e:
        # Try pdftotext if PyPDF2 fails
        try:
            with tempfile.NamedTemporaryFile(suffix='.txt') as temp_txt:
                subprocess.run(
                    ['pdftotext', str(file_path), temp_txt.name],
                    check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
                )
                with open(temp_txt.name, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        except Exception as pdf_e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}. pdftotext error: {str(pdf_e)}")


def extract_text_from_excel(file_path):
    """Extract text from Excel files directly"""
    try:
        import pandas as pd
        
        # Determine the engine based on file extension
        ext = file_path.suffix.lower()
        engine = 'xlrd' if ext == '.xls' else 'openpyxl'
        
        # Read all sheets
        all_sheets = []
        excel_file = pd.ExcelFile(file_path, engine=engine)
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            all_sheets.append(f"Sheet: {sheet_name}\n")
            all_sheets.append(df.to_string(index=False))
            all_sheets.append("\n\n")
        
        # Make sure to close the file to avoid resource warnings
        excel_file.close()
        return "".join(all_sheets)
    except ImportError:
        # Try to install pandas and related packages
        installed_pandas = install_dependency("pandas", "Installing pandas for Excel support...")
        
        if ext == '.xlsx':
            installed_openpyxl = install_dependency("openpyxl", "Installing openpyxl for XLSX support...")
        else:  # .xls
            installed_xlrd = install_dependency("xlrd", "Installing xlrd for XLS support...")
        
        if installed_pandas and ((ext == '.xlsx' and installed_openpyxl) or (ext == '.xls' and installed_xlrd)):
            # Try again after installation
            import pandas as pd
            excel_file = pd.ExcelFile(file_path, engine=engine)
            all_sheets = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                all_sheets.append(f"Sheet: {sheet_name}\n")
                all_sheets.append(df.to_string(index=False))
                all_sheets.append("\n\n")
            excel_file.close()
            return "".join(all_sheets)
        else:
            raise Exception(f"Failed to install Excel support. Please install manually: pip install pandas openpyxl xlrd")
    except Exception as e:
        raise Exception(f"Failed to extract text from Excel: {str(e)}")


def extract_text_from_csv(file_path):
    """Extract text from a CSV file directly with robust error handling"""
    try:
        # Try multiple parsing methods
        methods = [
            # Method 1: Standard pandas read_csv
            lambda: pd.read_csv(file_path, low_memory=False).to_string(index=False),
            
            # Method 2: Pandas with different delimiters
            lambda: pd.read_csv(file_path, sep=';', low_memory=False).to_string(index=False),
            
            # Method 3: Pandas with flexible parsing
            lambda: pd.read_csv(file_path, sep=None, engine='python', low_memory=False).to_string(index=False),
            
            # Method 4: Basic CSV reader with error handling
            lambda: _basic_csv_read(file_path)
        ]
        
        # Try each method in turn
        last_error = None
        for method in methods:
            try:
                return method()
            except Exception as e:
                last_error = e
                continue
        
        # If all methods failed, raise the last error
        raise last_error
        
    except ImportError:
        # Try to install pandas
        if install_dependency("pandas", "Installing pandas for CSV support..."):
            # Try again after installation
            import pandas as pd
            try:
                return pd.read_csv(file_path, low_memory=False).to_string(index=False)
            except Exception:
                return _basic_csv_read(file_path)
        else:
            return _basic_csv_read(file_path)
    except Exception as e:
        # Fall back to basic CSV reading
        try:
            return _basic_csv_read(file_path)
        except Exception as csv_e:
            raise Exception(f"Failed to extract text from CSV: {str(e)}, {str(csv_e)}")


def _basic_csv_read(file_path):
    """Basic CSV reader with robust error handling"""
    # Sniff the CSV dialect first
    try:
        with open(file_path, 'r', newline='', encoding='utf-8') as f:
            sample = f.read(4096)
            
        if not sample:
            return "Empty file"
            
        dialect = csv.Sniffer().sniff(sample)
        has_header = csv.Sniffer().has_header(sample)
        
        rows = []
        with open(file_path, 'r', newline='', encoding='utf-8') as f:
            reader = csv.reader(f, dialect)
            for row in reader:
                rows.append('\t'.join(row))
        
        return '\n'.join(rows)
        
    except Exception as e:
        # If sniffing fails, try a direct approach
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            # Last resort - try to read it as binary
            with open(file_path, 'rb') as f:
                data = f.read()
                try:
                    # Try UTF-8
                    return data.decode('utf-8')
                except UnicodeDecodeError:
                    # Try Latin-1
                    return data.decode('latin-1', errors='ignore')


def extract_text_from_xml(file_path):
    """Extract text from an XML file directly with robust error handling"""
    try:
        # Try reading as raw text first
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Replace XML/HTML tags with newlines
        import re
        text = re.sub(r'<[^>]*>', '\n', content)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        # Split into lines and filter empty lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        return "\n".join(lines)
    except Exception as raw_e:
        try:
            # Try to use lxml if available
            import lxml.etree as ET
            
            def get_text_from_element(element):
                """Recursively extract text from XML element and its children"""
                text = element.text or ""
                for child in element:
                    text += get_text_from_element(child)
                text += element.tail or ""
                return text
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            return get_text_from_element(root)
            
        except ImportError:
            # Try to install lxml
            if install_dependency("lxml", "Installing lxml for XML support..."):
                # Try again after installation
                import lxml.etree as ET
                tree = ET.parse(file_path)
                root = tree.getroot()
                def get_text_from_element(element):
                    text = element.text or ""
                    for child in element:
                        text += get_text_from_element(child)
                    text += element.tail or ""
                    return text
                return get_text_from_element(root)
            
        except Exception as lxml_e:
            # Last resort - try Python's built-in XML parser
            try:
                import xml.etree.ElementTree as ET
                
                tree = ET.parse(file_path)
                root = tree.getroot()
                
                text_parts = []
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        text_parts.append(elem.text.strip())
                
                return "\n".join(text_parts)
                
            except Exception as builtin_e:
                raise Exception(f"All XML parsing methods failed: {str(raw_e)}, {str(lxml_e)}, {str(builtin_e)}")


def try_unstructured_conversion(file_path):
    """Try to convert using Unstructured.io if available"""
    try:
        from unstructured.partition.auto import partition
        
        # Handle NLTK requirements for unstructured
        try:
            import nltk
            try:
                nltk.data.find('tokenizers/punkt')
            except LookupError:
                nltk.download('punkt')
        except (ImportError, Exception) as e:
            print(f"  NLTK error: {str(e)}")
            install_dependency("nltk", "Installing NLTK for text processing...")
            import nltk
            nltk.download('punkt')
        
        # Now try to partition the document
        elements = partition(str(file_path))
        text_content = "\n\n".join([str(element) for element in elements])
        return text_content
    except ImportError:
        raise Exception("Unstructured.io is not installed")
    except Exception as e:
        raise Exception(f"Unstructured.io conversion failed: {str(e)}")


def process_file(file_path, output_dir, supported_extensions, force=False, skip_unstructured=False):
    """Process a single file using appropriate method based on file type"""
    file_ext = file_path.suffix.lower()
    if file_ext not in supported_extensions:
        print(f"Skipping unsupported file: {file_path}")
        return False
    
    # Create output file path with .txt extension
    output_file = output_dir / f"{file_path.stem}.txt"
    
    # Check if output file already exists
    if output_file.exists() and not force:
        response = input(f"File {output_file} already exists. Overwrite? (y/n): ")
        if response.lower() not in ['y', 'yes']:
            print(f"Skipping {file_path}")
            return False
    
    print(f"Converting: {file_path}")
    
    text_content = None
    
    # Try Unstructured.io first if not disabled and not a known problematic format
    if not skip_unstructured and file_ext not in ['.doc', '.csv']:
        try:
            text_content = try_unstructured_conversion(file_path)
            print(f"  (Used Unstructured.io)")
        except Exception as e:
            print(f"  Unstructured.io failed: {str(e)}")
    
    # If Unstructured failed or not tried, use direct methods based on file type
    if text_content is None:
        try:
            if file_ext in ['.pdf']:
                text_content = extract_text_from_pdf(file_path)
                print(f"  (Used direct PDF extraction)")
            
            elif file_ext in ['.docx', '.doc']:
                if file_ext == '.docx':
                    text_content = extract_text_from_docx(file_path)
                else:
                    text_content = extract_text_from_doc(file_path)
                print(f"  (Used direct {file_ext} extraction)")
            
            elif file_ext in ['.xlsx', '.xls']:
                text_content = extract_text_from_excel(file_path)
                print(f"  (Used direct Excel extraction)")
            
            elif file_ext == '.csv':
                text_content = extract_text_from_csv(file_path)
                print(f"  (Used direct CSV extraction)")
            
            elif file_ext == '.xml':
                text_content = extract_text_from_xml(file_path)
                print(f"  (Used direct XML extraction)")
            
            elif file_ext in ['.txt', '.md', '.json', '.html', '.htm']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
                print(f"  (Used direct text file reading)")
            
            else:
                # Last resort - try to read as plain text
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text_content = f.read()
                    print(f"  (Used direct text file reading as fallback)")
                except Exception as txt_e:
                    raise Exception(f"No direct extraction method available for {file_ext}: {str(txt_e)}")
                
        except Exception as e:
            print(f"  Direct extraction failed: {str(e)}")
            return False
    
    # Write the text content
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text_content)
        print(f"✓ Successfully converted to: {output_file}")
        return True
    except Exception as e:
        print(f"✗ Error writing output file: {str(e)}")
        return False


def process_directory(directory, output_dir, supported_extensions, recursive=False, force=False, skip_unstructured=False):
    """Process all files in a directory"""
    print(f"Processing directory: {directory}")
    
    total_files = 0
    successful = 0
    failed = 0
    
    for item in directory.iterdir():
        if item.is_file():
            result = process_file(item, output_dir, supported_extensions, force, skip_unstructured)
            total_files += 1
            if result:
                successful += 1
            else:
                failed += 1
        elif item.is_dir() and recursive:
            # Create corresponding output subdirectory
            relative_path = item.relative_to(directory)
            new_output_dir = output_dir / relative_path
            new_output_dir.mkdir(exist_ok=True)
            
            # Process the subdirectory
            sub_total, sub_successful, sub_failed = process_directory(
                item, new_output_dir, supported_extensions, recursive, force, skip_unstructured
            )
            total_files += sub_total
            successful += sub_successful
            failed += sub_failed
    
    return total_files, successful, failed


def main():
    parser = argparse.ArgumentParser(description="All-in-One Document Converter with Dependency Fixing")
    parser.add_argument("input", help="Input file or directory path", nargs="?")
    parser.add_argument("-o", "--output", help="Output directory (default: ./text_output)")
    parser.add_argument("-r", "--recursive", action="store_true", help="Process directories recursively")
    parser.add_argument("--force", action="store_true", help="Overwrite existing text files without confirmation")
    parser.add_argument("--skip-unstructured", action="store_true", help="Skip trying Unstructured.io and use direct methods only")
    parser.add_argument("--check-deps", action="store_true", help="Check dependencies and exit")
    parser.add_argument("--install-deps", action="store_true", help="Install dependencies and exit")
    parser.add_argument("--install-nltk", action="store_true", help="Install NLTK data (punkt, etc.)")
    args = parser.parse_args()

    # Print welcome message
    print_section("All-in-One Document Converter")
    
    # Check/install dependencies if requested
    if args.check_deps or args.install_deps or args.install_nltk:
        check_and_install_dependencies(['pdf', 'docx', 'excel', 'csv', 'xml'])
        if args.install_nltk:
            try:
                import nltk
                print("Installing NLTK data...")
                nltk.download('punkt')
                nltk.download('averaged_perceptron_tagger')
                print("NLTK data installed successfully.")
            except Exception as e:
                print(f"Error installing NLTK data: {str(e)}")
        if args.check_deps:
            sys.exit(0)
    
    # Check if input is provided when needed
    if not args.input and not args.install_deps:
        parser.print_help()
        print("\nError: Input path is required")
        sys.exit(1)
    
    # Exit if only installing dependencies
    if args.install_deps:
        print("\nDependency installation complete. Run the script again with input parameters.")
        sys.exit(0)
    
    # Set up output directory
    output_dir = args.output if args.output else "./text_output"
    os.makedirs(output_dir, exist_ok=True)
    
    # Supported file extensions
    supported_extensions = [
        # Document formats
        '.pdf', '.pptx', '.ppt', '.docx', '.doc', '.xlsx', '.xls', '.rtf',
        # Image formats
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff',
        # Web and text formats
        '.html', '.htm', '.csv', '.json', '.xml', '.txt', '.md',
        # Archives
        '.zip'
    ]
    
    # Process files
    total_files = 0
    successful_conversions = 0
    failed_conversions = 0
    
    start_time = time.time()
    input_path = Path(args.input)
    
    # Check formats needed for this conversion
    required_formats = set()
    if input_path.is_file():
        ext = input_path.suffix.lower()
        if ext in ['.pdf']:
            required_formats.add('pdf')
        elif ext in ['.docx', '.doc']:
            required_formats.add('docx')
        elif ext in ['.xlsx', '.xls']:
            required_formats.add('excel')
        elif ext == '.csv':
            required_formats.add('csv')
        elif ext == '.xml':
            required_formats.add('xml')
    elif input_path.is_dir():
        # For directories, check the first few files
        found_formats = set()
        count = 0
        for item in input_path.rglob('*'):
            if count > 10:
                break
            if item.is_file():
                ext = item.suffix.lower()
                count += 1
                if ext in ['.pdf']:
                    required_formats.add('pdf')
                elif ext in ['.docx', '.doc']:
                    required_formats.add('docx')
                elif ext in ['.xlsx', '.xls']:
                    required_formats.add('excel')
                elif ext == '.csv':
                    required_formats.add('csv')
                elif ext == '.xml':
                    required_formats.add('xml')
    
    # Check and install dependencies if needed
    if required_formats:
        check_and_install_dependencies(required_formats)
    
    if input_path.is_file():
        result = process_file(input_path, Path(output_dir), supported_extensions, args.force, args.skip_unstructured)
        total_files = 1
        if result:
            successful_conversions = 1
        else:
            failed_conversions = 1
    elif input_path.is_dir():
        total_files, successful_conversions, failed_conversions = process_directory(
            input_path, Path(output_dir), supported_extensions, args.recursive, args.force, args.skip_unstructured
        )
    else:
        print(f"Error: Input path '{args.input}' does not exist")
        sys.exit(1)


    elapsed_time = time.time() - start_time
    print("\n" + "="*50)
    print("Conversion Summary:")
    print(f"Total files processed: {total_files}")
    print(f"Successful conversions: {successful_conversions}")
    print(f"Failed conversions: {failed_conversions}")
    print(f"Time elapsed: {elapsed_time:.2f} seconds")
    print("="*50)


if __name__ == "__main__":
    main()